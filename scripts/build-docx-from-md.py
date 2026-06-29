from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


def set_font(run, name: str, size: int | float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = col_widths[min(idx, len(col_widths) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0F172A")
    title.paragraph_format.space_after = Pt(11)

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_inline_markup(part) for part in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s:\-]+\|[\s:\-\|]+$", line.strip()))


def add_table(doc: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return
    header = split_table_row(table_lines[0])
    rows = [split_table_row(line) for line in table_lines[2:]]
    col_count = max(1, len(header))
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    widths = [9360 // col_count] * col_count
    widths[-1] += 9360 - sum(widths)

    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F2F4F7")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, "Calibri", 9, True, "0F172A")

    for row_values in rows:
        row = table.add_row()
        for idx in range(col_count):
            cell = row.cells[idx]
            cell.text = row_values[idx] if idx < len(row_values) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_font(run, "Calibri", 9, False, "1F2937")

    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    borders.append(bottom)
    p_pr.append(borders)


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        run = paragraph.add_run(line if line else " ")
        set_font(run, "Consolas", 8.5, False, "334155")
    doc.add_paragraph()


def add_markdown(doc: Document, path: Path, page_break_before: bool) -> None:
    if page_break_before:
        doc.add_page_break()

    lines = path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("<!--") and stripped.endswith("-->"):
            index += 1
            continue
        if stripped in {"<details>", "</details>"} or stripped.startswith("<summary>"):
            index += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = strip_inline_markup(heading.group(2))
            if level == 1:
                doc.add_paragraph(text, style="Title")
            elif level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(strip_inline_markup(bullet.group(1)))
            index += 1
            continue

        number = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(strip_inline_markup(number.group(1)))
            index += 1
            continue

        quote = re.match(r"^>\s*(.+)$", stripped)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        if quote:
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(strip_inline_markup(quote.group(1)))
            set_font(run, "Calibri", 10.5, False, "475569")
        else:
            paragraph.add_run(strip_inline_markup(stripped))
        index += 1


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a DOCX reading copy from Markdown files.")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--title", default="1Cati Documentation", help="Core document title")
    parser.add_argument("inputs", nargs="+", help="Markdown input files")
    args = parser.parse_args()

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = args.title
    doc.core_properties.author = "1Cati Product and Engineering"
    doc.core_properties.last_modified_by = "Codex Documents workflow"
    doc.core_properties.keywords = "1Cati, Ataberk Estate, CRM, Supabase, requirements"

    for idx, input_path in enumerate(args.inputs):
        path = (ROOT / input_path).resolve() if not Path(input_path).is_absolute() else Path(input_path)
        add_markdown(doc, path, page_break_before=idx > 0)

    output = (ROOT / args.output).resolve() if not Path(args.output).is_absolute() else Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
